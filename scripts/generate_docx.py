"""
Generates Project_Explanation.docx — a plain-English walkthrough of the
AI Agent Ticket Triage & Auto-Routing System.

Run:
    python scripts/generate_docx.py
Output:
    Project_Explanation.docx  (project root)
"""
from __future__ import annotations
import pathlib
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Colour helpers ────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x00, 0x15, 0x29)
BLUE   = RGBColor(0x16, 0x77, 0xFF)
GREEN  = RGBColor(0x52, 0xC4, 0x1A)
ORANGE = RGBColor(0xFA, 0x8C, 0x16)
RED    = RGBColor(0xF5, 0x22, 0x2D)
PURPLE = RGBColor(0x72, 0x2E, 0xD1)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
GREY   = RGBColor(0x59, 0x59, 0x59)
LIGHT  = RGBColor(0xF0, 0xF2, 0xF5)


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def cell_text(cell, text, bold=False, size=10, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


# ── Document setup ────────────────────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Default style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def h1(text, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = color
    run.font.name = "Calibri"
    return p


def h2(text, color=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = color
    run.font.name = "Calibri"
    return p


def h3(text, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = color
    run.font.name = "Calibri"
    return p


def body(text, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.italic = italic
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color
    return p


def bullet(text, level=0, color=None, bold_prefix: str = ""):
    style_name = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style_name)
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        run = p.add_run(bold_prefix + "  ")
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Calibri"
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color
    return p


def divider():
    p = doc.add_paragraph("─" * 90)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0xD9, 0xD9, 0xD9)
        run.font.size = Pt(8)


def coloured_box_para(text, bg_hex, text_color=WHITE):
    """Simulate a coloured header using a 1×1 table."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.rows[0].cells[0]
    cell.width = Inches(6.5)
    set_cell_bg(cell, bg_hex)
    cell_text(cell, text, bold=True, size=12, color=text_color,
              align=WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_paragraph()  # spacer


# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_para.paragraph_format.space_before = Pt(30)
t_run = title_para.add_run("AI Agent Ticket Triage &\nAuto-Routing System")
t_run.bold = True
t_run.font.size = Pt(26)
t_run.font.color.rgb = NAVY
t_run.font.name = "Calibri"

sub_para = doc.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
s_run = sub_para.add_run("Complete Project Explanation — Workflow, Components & Logic")
s_run.font.size = Pt(13)
s_run.font.color.rgb = GREY
s_run.italic = True
s_run.font.name = "Calibri"

date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
d_run = date_para.add_run("June 2026")
d_run.font.size = Pt(11)
d_run.font.color.rgb = GREY
d_run.font.name = "Calibri"

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — WHAT IS THIS PROJECT?
# ══════════════════════════════════════════════════════════════════════════════
h1("1.  What Is This Project?")
divider()
body(
    "The AI Agent Ticket Triage & Auto-Routing System is an intelligent software solution "
    "that automates the process of handling IT support tickets. Instead of a human analyst "
    "manually reading each ticket, deciding what type of problem it is, how urgent it is, "
    "and who should fix it — this system does all of that automatically using Artificial Intelligence."
)
body(
    "The system reads tickets from a standard Excel (.xlsx) file, runs them through a pipeline "
    "of five specialised AI agents, and produces an enriched output file with routing decisions — "
    "all within seconds."
)

h2("1.1  The Problem It Solves")
tbl = doc.add_table(rows=2, cols=2)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
set_cell_bg(tbl.rows[0].cells[0], "F5222D")
set_cell_bg(tbl.rows[0].cells[1], "52C41A")
cell_text(tbl.rows[0].cells[0], "BEFORE  —  Manual Process", bold=True, size=11, color=WHITE)
cell_text(tbl.rows[0].cells[1], "AFTER  —  AI Agent System",  bold=True, size=11, color=WHITE)

before_lines = (
    "• 15–30 minutes to triage and route a batch of tickets\n"
    "• Inconsistent — different analysts make different calls\n"
    "• No priority intelligence — critical issues get buried\n"
    "• Zero visibility into routing trends or team workload"
)
after_lines = (
    "• Under 2 minutes for a full batch — fully automated\n"
    "• Consistent — same AI logic applied to every ticket\n"
    "• Automatic P1–P4 priority assignment with reasoning\n"
    "• Real-time KPI dashboard, charts, and team metrics"
)
cell_text(tbl.rows[1].cells[0], before_lines, size=10)
cell_text(tbl.rows[1].cells[1], after_lines,  size=10)
doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — HOW TO RUN THE PROJECT
# ══════════════════════════════════════════════════════════════════════════════
h1("2.  How to Run the Project")
divider()

h2("2.1  One-Time Setup  (do this once)")
body("Open a terminal in the project folder and run these commands in order:")

steps_setup = [
    ("pip install -e .",
     "Installs all Python libraries and makes the project importable."),
    ("python src/ml/trainer.py",
     "Trains the Machine Learning model on the sample data. This MUST be done "
     "before uploading tickets — otherwise every ticket falls back to the LLM, "
     "making processing very slow."),
    ("python scripts/generate_routing_config.py",
     "Creates data/routing_config.xlsx — the file that maps categories to "
     "teams, team leads, and member email addresses."),
]
for cmd, desc in steps_setup:
    p = doc.add_paragraph(style="List Number")
    r1 = p.add_run(cmd + "\n")
    r1.bold = True
    r1.font.name = "Courier New"
    r1.font.size = Pt(10)
    r1.font.color.rgb = NAVY
    r2 = p.add_run("    " + desc)
    r2.font.size = Pt(10)
    r2.font.name = "Calibri"

h2("2.2  Every Time You Run the Project  (2 terminals)")

coloured_box_para("Terminal 1  —  Start the Backend API", "001529")
p = doc.add_paragraph()
r = p.add_run("    uvicorn src.api.main:app --reload")
r.bold = True; r.font.name = "Courier New"; r.font.size = Pt(11); r.font.color.rgb = BLUE
body("    This starts the Python server on http://localhost:8000.  All AI processing "
     "happens here.")

coloured_box_para("Terminal 2  —  Start the Frontend UI", "1677FF")
p = doc.add_paragraph()
r = p.add_run("    cd frontend  →  npm run dev")
r.bold = True; r.font.name = "Courier New"; r.font.size = Pt(11); r.font.color.rgb = WHITE
body("    This starts the React web application on http://localhost:5173.  This is what "
     "the user sees in the browser.")

h2("2.3  Environment Variables  (.env file)")
body("Create a .env file in the project root.  The most important settings are:")

env_rows = [
    ("OLLAMA_MODEL",           "gpt-oss:120b-cloud",          "Which LLM to use for AI decisions"),
    ("OLLAMA_BASE_URL",        "http://localhost:11434",       "Where the Ollama LLM server is running"),
    ("EMAIL_USERNAME",         "you@gmail.com",                "Gmail address for sending email alerts"),
    ("EMAIL_PASSWORD",         "your-app-password",            "Gmail App Password (not your login password)"),
    ("CONFIDENCE_AUTO_ROUTE",  "0.72",                         "Score above which a ticket is auto-routed"),
    ("CONFIDENCE_FLAG",        "0.55",                         "Score above which a ticket is flagged for review"),
    ("DATABASE_URL",           "sqlite+aiosqlite:///./tickets.db", "Local SQLite database file"),
]
tbl = doc.add_table(rows=1 + len(env_rows), cols=3)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
headers = ["Variable", "Default Value", "What It Controls"]
for i, h in enumerate(headers):
    set_cell_bg(tbl.rows[0].cells[i], "001529")
    cell_text(tbl.rows[0].cells[i], h, bold=True, size=10, color=WHITE)
for ri, (var, val, desc) in enumerate(env_rows):
    row = tbl.rows[ri + 1]
    cell_text(row.cells[0], var,  bold=True,  size=9, color=NAVY)
    cell_text(row.cells[1], val,  bold=False, size=9)
    cell_text(row.cells[2], desc, bold=False, size=9)
doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — END-TO-END WORKFLOW
# ══════════════════════════════════════════════════════════════════════════════
h1("3.  End-to-End Workflow — What Happens When You Upload an Excel File?")
divider()
body(
    "This is the core of the system.  When you upload a tickets.xlsx file, the following "
    "six stages happen automatically in sequence."
)

# Stage boxes
stages = [
    ("001529", "Stage 1",  "Excel File Upload",
     "The user opens the web portal at http://localhost:5173 and drags the Excel file "
     "onto the Upload page.  The file is sent to the FastAPI backend as raw bytes.  "
     "The system reads every row using the pandas library.  If any row is missing a "
     "Ticket_ID, one is automatically generated using the current timestamp "
     "(e.g., TKT-20260625143022-0001) so that each upload is unique and will never "
     "collide with a previous upload."),
    ("1677FF", "Stage 2",  "Per-Ticket AI Processing",
     "Each ticket is independently passed through a pipeline of five AI agents.  "
     "Agents 1 and 2 (Classifier and Priority) run at the same time in parallel to "
     "save processing time.  After they finish, Agents 3 and 4 run in sequence.  "
     "This entire pipeline is described in detail in Section 4."),
    ("722ED1", "Stage 3",  "Confidence Decision",
     "After all agents have run, a final confidence score is calculated for each ticket. "
     "This score determines the routing outcome — see Section 5 for the exact formula."),
    ("52C41A", "Stage 4",  "Email Notification",
     "If a ticket scores above the auto-route threshold (≥ 0.72 by default), an HTML "
     "email is immediately sent via Gmail SMTP to the assigned team member.  The email "
     "includes the ticket summary, AI-predicted category, priority, and the assignee's name."),
    ("FA8C16", "Stage 5",  "Database Storage",
     "All processed tickets — including both the original values and the AI predictions — "
     "are saved to a SQLite database file (tickets.db).  Duplicate ticket IDs are "
     "automatically detected and skipped so re-uploading the same file is safe."),
    ("001529", "Stage 6",  "Enriched Excel Download",
     "The user receives a download of the original Excel file with 8 new AI columns added "
     "side-by-side with the original columns for easy comparison:\n"
     "    Original_Category | AI_Category | Original_Priority | AI_Priority | "
     "Original_Team | AI_Assigned_Team | Final_Confidence | Routing_Status | Email_Sent"),
]

for bg, stage, title, desc in stages:
    coloured_box_para(f"{stage}  —  {title}", bg)
    p = doc.add_paragraph(desc)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(8)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = "Calibri"

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — THE FIVE AI AGENTS
# ══════════════════════════════════════════════════════════════════════════════
h1("4.  The Five AI Agents — What Each One Does")
divider()
body(
    "The system uses five specialised agents.  Each agent has one job and one job only.  "
    "They work together like an assembly line — each one picks up where the previous one left off."
)

agents = [
    (BLUE,   "Agent 1 — Classifier Agent",
     "File: src/agents/classifier_agent.py",
     "Question it answers: What TYPE of problem is this ticket?",
     [
         "The agent reads the ticket's Summary and Description fields.",
         "It first tries the Machine Learning (ML) model (a TF-IDF + Logistic Regression "
          "classifier trained on sample tickets).  If the ML model's confidence is 55% or "
          "above, its answer is used.",
         "If the ML model is less than 55% confident (which happens when the ticket is "
          "unusual or the model hasn't seen enough similar examples), the agent calls "
          "Ollama — the local Large Language Model — as a fallback.  Ollama reads the "
          "ticket in plain English and makes its own judgement.",
         "Output: a category string such as 'Database connection timeout', "
          "'Phishing email reported', 'VPN connectivity issue', etc., along with "
          "a confidence score between 0.0 and 1.0.",
     ]),
    (ORANGE, "Agent 2 — Priority Agent",
     "File: src/agents/priority_agent.py",
     "Question it answers: How URGENT is this ticket?  (P1 / P2 / P3 / P4)",
     [
         "The agent scans the ticket text for keyword signals first (fast path).",
         "Words like 'down', 'outage', 'breach', 'ransomware', 'all users affected' "
          "immediately trigger P1 (Critical).",
         "Words like 'degraded', 'slow', 'multiple users affected' trigger P2 (High).",
         "Words like 'cosmetic', 'enhancement', 'documentation' trigger P4 (Low).",
         "If no keywords are matched, Ollama is asked to assess the urgency based on "
          "the full ticket text and business context.  Ollama returns P1–P4 with a "
          "confidence score and a short reasoning note.",
         "P3 (Medium) is the default when there is no clear signal.",
         "This agent runs at the same time as Agent 1 (parallel execution) to save time.",
     ]),
    (GREEN,  "Agent 3 — Routing Agent",
     "File: src/agents/routing_agent.py",
     "Question it answers: WHO should handle this ticket?",
     [
         "The agent reads data/routing_config.xlsx — a spreadsheet that lists every team, "
          "team lead, team members, their email addresses, and their roles.",
         "It uses a built-in mapping table to link each category to a team "
          "(e.g., 'Database connection timeout' → Database Team).",
         "For P1 and P2 tickets, the ticket always goes directly to the Team Lead — "
          "because high-priority issues need senior attention immediately.",
         "For P3 and P4 tickets, Ollama selects the most appropriate team member "
          "based on the ticket description and each member's role.",
         "Output: team name, assignee name, assignee email, and team lead email.",
         "To add a new team or change assignments — simply update routing_config.xlsx.  "
          "No code changes are needed.",
     ]),
    (NAVY,   "Agent 4 — Confidence Agent",
     "File: src/agents/confidence_agent.py",
     "Question it answers: Are we CONFIDENT ENOUGH to act automatically?",
     [
         "This agent does not call any AI model.  It is pure maths.",
         "It combines the outputs of the previous agents into a single confidence score "
          "between 0.0 and 1.0.",
         "Two calculation paths exist depending on which classifier was used:",
         "    PATH A — If the LLM was used (ML was below 55% threshold):\n"
          "        Score  =  LLM confidence × 0.70  +  Priority confidence × 0.30\n"
          "        Reason: The low ML score that triggered the fallback would drag the "
          "result down if included.  The LLM is the reliable signal here.",
         "    PATH B — If the ML model was used alone:\n"
          "        Score  =  ML confidence × 0.80  +  Priority confidence × 0.20\n"
          "        Reason: The ML model was confident, so it gets the dominant weight.",
         "The final score is then compared to the configured thresholds to decide the routing outcome (see Section 5).",
     ]),
    (PURPLE, "Agent 5 — Feedback Agent",
     "File: src/agents/feedback_agent.py",
     "Question it answers: How do we get SMARTER over time?",
     [
         "This agent activates when a human reviewer approves or corrects a ticket "
          "from the Review Queue in the web portal.",
         "It saves the human's correction — what the correct category, priority, and "
          "team should have been — to the database.",
         "When a reviewer submits a correction, the ticket's routing status is "
          "automatically updated to 'Manual' and it disappears from the Review Queue.",
         "One correction per ticket is allowed.  Submitting the same ticket twice is "
          "blocked with an error message.",
         "The Human Corrections counter on the Dashboard counts unique tickets "
          "corrected — not total submissions — so it cannot inflate artificially.",
         "When the total corrections reach a configurable threshold (default: 20), "
          "a log message prompts the team to retrain the ML model with the new "
          "correction data, which improves future accuracy.",
     ]),
]

for color, title, filepath, question, points in agents:
    h2(title, color=color)
    body(filepath, italic=True, color=GREY)
    p = doc.add_paragraph()
    r = p.add_run(question)
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = color; r.font.name = "Calibri"
    for pt in points:
        bullet(pt)
    doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — CONFIDENCE THRESHOLDS & ROUTING DECISIONS
# ══════════════════════════════════════════════════════════════════════════════
h1("5.  Confidence Thresholds — How Routing Decisions Are Made")
divider()
body(
    "After Agent 4 calculates the final confidence score, it is compared to two "
    "configurable thresholds to decide what happens to the ticket."
)

tbl = doc.add_table(rows=4, cols=4)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

# Header row
for i, h in enumerate(["Score Range", "Decision", "What Happens", "Dashboard Count"]):
    set_cell_bg(tbl.rows[0].cells[i], "001529")
    cell_text(tbl.rows[0].cells[i], h, bold=True, size=10, color=WHITE)

rows_data = [
    ("52C41A", WHITE, "≥ 0.72  (configurable)",
     "AUTO-ROUTED",
     "Ticket is assigned to the correct team member.  An HTML email is sent "
     "immediately to the assignee via Gmail SMTP.  No human action required.",
     "Auto-Routed count ↑"),
    ("FA8C16", WHITE, "0.55 – 0.72  (configurable)",
     "FLAGGED",
     "AI is not confident enough to act alone.  Ticket appears in the Review "
     "Queue.  A human reviewer can approve the AI suggestion or override the "
     "team/priority before approving.",
     "Flagged for Review count ↑"),
    ("F5222D", WHITE, "< 0.55",
     "HELD",
     "The AI prediction is too uncertain.  Ticket is held for full manual "
     "classification.  Appears in Review Queue with a red highlight.",
     "Held count ↑"),
]

for ri, (bg, txt_color, score, decision, description, dashboard) in enumerate(rows_data):
    row = tbl.rows[ri + 1]
    set_cell_bg(row.cells[0], bg)
    set_cell_bg(row.cells[1], bg)
    cell_text(row.cells[0], score,      bold=True, size=10, color=txt_color)
    cell_text(row.cells[1], decision,   bold=True, size=10, color=txt_color)
    cell_text(row.cells[2], description, size=9)
    cell_text(row.cells[3], dashboard,  size=9)

doc.add_paragraph()
body(
    "Important: These thresholds can be changed at any time by editing the "
    ".env file.  No code changes are needed.  As the ML model improves through "
    "feedback, the thresholds can be raised to require stricter confidence "
    "before auto-routing.",
    italic=True
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — THE WEB PORTAL (FRONTEND)
# ══════════════════════════════════════════════════════════════════════════════
h1("6.  The Web Portal — 5 Pages Explained")
divider()
body("The frontend is a React web application styled with Ant Design.  "
     "It runs at http://localhost:5173 and communicates with the backend API.")

pages = [
    (BLUE,   "Dashboard",
     "The first page you see.  Displays real-time KPI cards: total tickets "
     "processed, auto-route percentage, flagged count, held count, average AI "
     "confidence, model accuracy, and human corrections.  Also shows a bar chart "
     "of ticket volume by team and a pie chart of priority distribution.  "
     "All numbers update automatically from the database."),
    (ORANGE, "Upload Tickets",
     "Drag-and-drop area for uploading the Excel file.  After upload, an animated "
     "progress bar shows the four processing steps: Uploading → Processing → "
     "Enriching → Complete.  When done, a Download button appears to save the "
     "enriched Excel file containing all AI predictions alongside original values."),
    (NAVY,   "All Tickets",
     "A paginated table showing every ticket stored in the database.  Columns "
     "include Ticket ID, summary, AI category, AI priority (colour-coded: P1 red, "
     "P2 orange, P3 blue, P4 green), assigned team, confidence score, routing "
     "status, and whether an email was sent.  Filter by team or routing status."),
    (ORANGE, "Review Queue",
     "Shows only tickets with status 'Flagged' or 'Held' — the ones that need "
     "human attention.  For each ticket, the reviewer can optionally override the "
     "AI-suggested team or priority using dropdown selectors.  Clicking Approve "
     "submits the correction, immediately removes the ticket from the queue, and "
     "updates the Dashboard counts.  Each ticket can only be reviewed once — "
     "the Approve button is disabled after submission to prevent duplicate entries."),
    (GREEN,  "Model Metrics",
     "Shows the performance statistics of the ML model: accuracy score, F1-macro "
     "score (which measures how well the model handles all categories equally), "
     "training sample count, and model version.  Also shows KPI vs target "
     "comparisons and a routing outcome breakdown.  Useful for monitoring model "
     "quality and deciding when to retrain."),
]

for color, title, desc in pages:
    h3(f"  {title}", color=color)
    p = doc.add_paragraph(desc)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(11); run.font.name = "Calibri"

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — TECHNOLOGY STACK
# ══════════════════════════════════════════════════════════════════════════════
h1("7.  Technology Stack — What Each Tool Does")
divider()

tech_rows = [
    ("Python 3.13",             "Core programming language for the entire backend."),
    ("FastAPI",                 "Web framework that exposes the REST API.  Handles file uploads, "
                                "ticket listing, feedback submission, and metrics."),
    ("SQLAlchemy + aiosqlite",  "Database library.  Stores tickets, feedback corrections, and "
                                "model metrics in SQLite locally.  Can be switched to Azure "
                                "PostgreSQL by changing one environment variable (DATABASE_URL)."),
    ("scikit-learn",            "Machine learning library.  Runs TF-IDF text vectorisation and "
                                "Logistic Regression classification.  The trained model is saved "
                                "to src/ml/models/classifier.pkl and loaded at startup."),
    ("Ollama (LLM)",            "Local large language model server.  Model: gpt-oss:120b-cloud.  "
                                "Used as fallback classifier when ML confidence is low, and as "
                                "the primary agent for priority assessment and member selection."),
    ("pandas + openpyxl",       "Read and write Excel files.  pandas handles data transformation; "
                                "openpyxl adds colour-coded priority formatting to the output file."),
    ("React 18 + Vite",         "Frontend JavaScript framework.  Single-page application that "
                                "runs in the browser and talks to the FastAPI backend."),
    ("Ant Design 5",            "Enterprise UI component library.  Provides tables, cards, "
                                "upload components, step indicators, and the sidebar layout."),
    ("Recharts",                "Chart library for React.  Used for the team bar chart and "
                                "priority pie chart on the Dashboard."),
    ("Gmail SMTP (smtplib)",    "Python's built-in email library.  Sends HTML notification "
                                "emails to assignees when a ticket is auto-routed.  Uses Gmail "
                                "App Password authentication — NOT the account password."),
    ("Pydantic + pydantic-settings",
                                "Data validation and environment configuration.  All settings "
                                "are loaded from the .env file through the Settings class."),
    ("Prometheus",              "Metrics collection.  The backend exposes a /prometheus-metrics "
                                "endpoint.  Optional — used for monitoring in production."),
]

tbl = doc.add_table(rows=1 + len(tech_rows), cols=2)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, h in enumerate(["Technology", "What It Does in This Project"]):
    set_cell_bg(tbl.rows[0].cells[i], "001529")
    cell_text(tbl.rows[0].cells[i], h, bold=True, size=10, color=WHITE)
for ri, (tech, desc) in enumerate(tech_rows):
    row = tbl.rows[ri + 1]
    cell_text(row.cells[0], tech, bold=True, size=9, color=NAVY)
    cell_text(row.cells[1], desc, size=9)
doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — KRA ACHIEVEMENTS
# ══════════════════════════════════════════════════════════════════════════════
h1("8.  KRA Achievements & Current Status")
divider()

kras = [
    (GREEN,  "✅  COMPLETE",
     "Data Collection & Synthetic Dataset",
     "50 labeled sample tickets generated across 7 teams (Testing, EDI, BizTalk, "
     "Database, Network, Access, Security).  4 priority levels.  Used to train the "
     "ML model.  Baseline accuracy metrics established."),
    (GREEN,  "✅  COMPLETE",
     "ML Classifier — 70–80% Baseline Accuracy",
     "TF-IDF vectoriser (bigram, 10,000 features) combined with Logistic Regression "
     "(C=5.0, multinomial) achieves 70–80% accuracy on the training data.  "
     "Accuracy, F1-macro, and confusion matrix are recorded after each training run."),
    (GREEN,  "✅  COMPLETE",
     "Multi-Agent Architecture",
     "Five specialised agents built and connected: Classifier, Priority, Routing, "
     "Confidence, and Feedback.  Agents run in an async pipeline per ticket."),
    (GREEN,  "✅  COMPLETE",
     "Excel Auto-Routing Pipeline",
     "Full upload-to-download pipeline working.  Confidence thresholds correctly "
     "route tickets to Auto-Routed / Flagged / Held.  Enriched Excel includes "
     "original and AI columns side-by-side for easy comparison."),
    (GREEN,  "✅  COMPLETE",
     "Real Cloud Action — Gmail SMTP Email Notification",
     "Live HTML email sent to the assignee's inbox for every auto-routed ticket.  "
     "Team lead digest emails sent for flagged batches.  Uses Gmail App Password "
     "authentication.  This is the required real-cloud-operation KRA item."),
    (ORANGE, "🔄  IN PROGRESS",
     "Feedback Loop & Continuous Improvement",
     "Human corrections are stored in the database.  Retraining is triggered when "
     "the correction count reaches the threshold.  The ML model accuracy is expected "
     "to reach ≥85% target after sufficient feedback data is collected and used for "
     "retraining."),
]

for color, status, title, desc in kras:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    r1 = p.add_run(status + "  ")
    r1.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = color; r1.font.name = "Calibri"
    r2 = p.add_run(title)
    r2.bold = True; r2.font.size = Pt(11); r2.font.color.rgb = NAVY; r2.font.name = "Calibri"
    p2 = doc.add_paragraph(desc)
    p2.paragraph_format.left_indent = Inches(0.4)
    p2.paragraph_format.space_after = Pt(4)
    for run in p2.runs:
        run.font.size = Pt(10); run.font.name = "Calibri"

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — QUICK REFERENCE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
h1("9.  Quick Reference — The Entire System in One Page")
divider()

body("Use this page when explaining the project to someone for the first time.", italic=True)

h2("One-sentence summary")
body(
    '"An Excel file of IT tickets goes in — the AI reads each one, decides what type of problem '
    'it is, how urgent it is, and who should fix it, then sends an email to that person '
    'automatically.  Anything the AI is unsure about goes to a human review queue instead."'
)

h2("The flow in 6 steps")
six_steps = [
    "User uploads tickets.xlsx through the web portal.",
    "The system reads every ticket row from the Excel file.",
    "Five AI agents process each ticket: classify → prioritise → route → score → record.",
    "The confidence score is checked against two thresholds.",
    "High confidence (≥0.72) → email sent automatically.  "
     "Medium (0.55–0.72) → goes to review queue.  "
     "Low (<0.55) → held for manual handling.",
    "User downloads the enriched Excel file with all AI decisions added as new columns.",
]
for i, step in enumerate(six_steps, 1):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(step)
    r.font.size = Pt(11); r.font.name = "Calibri"

h2("Three numbers to remember")
key_numbers = [
    ("0.72",  "The auto-route threshold.  Tickets scoring above this are handled without human involvement."),
    ("0.55",  "The flag threshold.  Tickets between 0.55 and 0.72 need a human to confirm before routing."),
    ("80/20", "The confidence formula weight split — 80% classifier (ML or LLM), 20% priority signal."),
]
for num, desc in key_numbers:
    p = doc.add_paragraph()
    r1 = p.add_run(f"  {num}  — ")
    r1.bold = True; r1.font.size = Pt(13); r1.font.color.rgb = BLUE; r1.font.name = "Calibri"
    r2 = p.add_run(desc)
    r2.font.size = Pt(11); r2.font.name = "Calibri"

h2("Seven teams supported")
body(
    "Testing Team  •  EDI Team  •  BizTalk Team  •  Database Team  "
    "•  Network Team  •  Access Team  •  Security Team\n"
    "To add a new team: update data/routing_config.xlsx.  No code changes required."
)

h2("The feedback loop")
body(
    "Every time a human reviewer corrects an AI decision, that correction is saved.  "
    "When 20 corrections have been collected, the system logs a prompt to retrain the "
    "ML model.  After retraining, the model is more accurate and will auto-route a "
    "higher percentage of tickets correctly."
)

# ── Footer ────────────────────────────────────────────────────────────────────
doc.add_paragraph()
divider()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    "AI Agent Ticket Triage & Auto-Routing System  |  "
    "Python · FastAPI · React · Ollama · scikit-learn · SQLite · Gmail SMTP"
)
r.font.size = Pt(9); r.font.color.rgb = GREY; r.font.name = "Calibri"

# ── Save ──────────────────────────────────────────────────────────────────────
out = pathlib.Path(__file__).resolve().parent.parent / "Project_Explanation.docx"
doc.save(str(out))
print(f"Saved -> {out}")
